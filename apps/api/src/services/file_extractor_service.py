import io
import logging
from typing import Tuple, Dict, Any, List
import chardet
import pandas as pd
import pdfplumber
import docx

logger = logging.getLogger("file_extractor_service")

MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024  # 10MB limit

class FileExtractionError(Exception):
    pass

def validate_magic_bytes_and_extension(filename: str, content_bytes: bytes) -> str:
    if len(content_bytes) > MAX_FILE_SIZE_BYTES:
        raise FileExtractionError(f"File size ({len(content_bytes)} bytes) exceeds maximum allowed limit of 10MB.")

    ext = filename.split(".")[-1].lower() if "." in filename else ""
    allowed_exts = ("pdf", "docx", "csv", "txt", "json", "md", "png", "jpg", "jpeg", "webp", "svg")
    if ext not in allowed_exts:
        raise FileExtractionError(f"Unsupported file extension '.{ext}'. Allowed types: PDF, DOCX, CSV, TXT, JSON, MD, PNG, JPG, JPEG, WEBP, SVG.")

    # Magic Bytes Validation
    if ext == "pdf":
        if not content_bytes.startswith(b"%PDF-"):
            raise FileExtractionError("Invalid PDF file: Magic bytes mismatch (%PDF- header missing).")
    elif ext == "docx":
        if not content_bytes.startswith(b"PK\x03\x04"):
            raise FileExtractionError("Invalid DOCX file: Magic bytes mismatch (Zip header missing).")

    return ext

def parse_pdf_with_page_info(content_bytes: bytes) -> str:
    """Extract PDF text preserving page numbers and table structures."""
    page_blocks = []
    with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            tables_text = ""
            try:
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 1:
                        header = table[0]
                        rows_formatted = []
                        for row in table[1:]:
                            cells = [f"{header[i]}: {row[i]}" for i in range(min(len(header), len(row))) if row[i]]
                            if cells:
                                rows_formatted.append(" | ".join(cells))
                        if rows_formatted:
                            tables_text += "\n" + "\n".join(rows_formatted)
            except Exception:
                pass

            combined_page = (text.strip() + "\n" + tables_text.strip()).strip()
            if combined_page:
                page_blocks.append(f"--- Page {idx} ---\n{combined_page}")

    return "\n\n".join(page_blocks)

def parse_docx_with_structure(content_bytes: bytes) -> str:
    """Extract DOCX text preserving paragraph headings and table structures."""
    doc = docx.Document(io.BytesIO(content_bytes))
    structured_parts = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        style_name = (p.style.name or "").lower()
        if "heading" in style_name:
            structured_parts.append(f"## {txt}")
        else:
            structured_parts.append(txt)

    for idx, table in enumerate(doc.tables, start=1):
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        table_lines = [f"=== Table {idx} ==="]
        for r_idx, row in enumerate(table.rows[1:], start=1):
            row_vals = [cell.text.strip() for cell in row.cells]
            paired = [f"{headers[i]}: {row_vals[i]}" for i in range(min(len(headers), len(row_vals))) if row_vals[i]]
            if paired:
                table_lines.append(f"Row {r_idx}: " + " | ".join(paired))
        if len(table_lines) > 1:
            structured_parts.append("\n".join(table_lines))

    return "\n\n".join(structured_parts)

def parse_csv_to_meaningful_records(content_bytes: bytes) -> str:
    """Converts CSV rows into structured key-value records instead of raw CSV line concatenation."""
    detected = chardet.detect(content_bytes)
    encoding = detected.get("encoding") or "utf-8"
    
    try:
        df = pd.read_csv(io.BytesIO(content_bytes), encoding=encoding)
    except Exception:
        df = pd.read_csv(io.BytesIO(content_bytes), encoding="utf-8", on_bad_lines="skip")

    df = df.fillna("")
    cols = df.columns.tolist()
    records = []

    for idx, row in df.iterrows():
        row_str = " | ".join([f"{col}: {row[col]}" for col in cols if str(row[col]).strip() != ""])
        if row_str:
            records.append(f"[Record {idx + 1}] {row_str}")

    return "\n\n".join(records)

def extract_text_from_file(filename: str, content_bytes: bytes) -> Tuple[str, int]:
    """
    Main file parsing dispatcher for PDF, DOCX, CSV, Markdown, and TXT files.
    """
    ext = validate_magic_bytes_and_extension(filename, content_bytes)
    extracted_text = ""

    try:
        if ext == "pdf":
            extracted_text = parse_pdf_with_page_info(content_bytes)
        elif ext == "docx":
            extracted_text = parse_docx_with_structure(content_bytes)
        elif ext == "csv":
            extracted_text = parse_csv_to_meaningful_records(content_bytes)
        elif ext in ("png", "jpg", "jpeg", "webp", "svg"):
            extracted_text = f"Image Media File: {filename}"
        elif ext in ("txt", "json", "md"):
            detected = chardet.detect(content_bytes)
            encoding = detected.get("encoding") or "utf-8"
            try:
                extracted_text = content_bytes.decode(encoding)
            except Exception:
                extracted_text = content_bytes.decode("utf-8", errors="ignore")

    except FileExtractionError:
        raise
    except Exception as e:
        logger.error(f"Text extraction failed for {filename}: {e}", exc_info=True)
        raise FileExtractionError(f"Failed to extract text from {filename}: {str(e)}")

    extracted_text_clean = extracted_text.strip()
    if not extracted_text_clean:
        extracted_text_clean = f"Document: {filename}"

    return extracted_text_clean, len(content_bytes)
